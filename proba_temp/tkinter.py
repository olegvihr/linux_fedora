from tkinter import *
from docx import Document
from docx.shared import Inches


def save_doc():
    # Создание нового документа Word
    doc = Document()

    # Добавление заголовка
    doc.add_heading('Тестовый документ', 0)

    # Добавление текста
    doc.add_paragraph(text_box.get(1.0, END))

    # Сохранение документа
    doc.save('test.docx')

# Создание окна Tkinter
root = Tk()

# Добавление элементов управления
label = Label(root, text="Введите текст:")
label.pack()

text_box = Text(root)
text_box.pack()

button = Button(root, text="Сохранить", command=save_doc)
button.pack()

# Запуск главного цикла окна Tkinter
root.mainloop()
