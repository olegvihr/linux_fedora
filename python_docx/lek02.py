from docx import Document
from docx.shared import Inches

guestnames = ["Иван", "Олег", "Ольга", "Мария и Александр"]

for guest in guestnames:
    document = Document()
    document.add_heading('Приглашение на свадьбу', 0)
    p = document.add_paragraph(guest + ', спешим пригласить вас на свадьбу')
    p.add_run('Анастасии и Дмитрия Макаренковых').bold = True
    p.add_run(', которая состоится 12.05.2020 в усадбе Середниково.')
    document.add_heading('Ждем вас на наш праздник!', level=1)
    document.add_paragraph('')
    document.add_picture('7svheiDEUx8.jpg', width=Inches(5.25))
    document.add_page_break()
    document.save('Приглашения/ ' + guest + '.docx')
