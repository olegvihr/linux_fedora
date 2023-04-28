from openpyxl import load_workbook
from docx import Document

# открыть файл Excel и лист
wb = load_workbook('file.xlsx')
ws = wb.active

# выбрать строку и столбец для данных
date = ws['A2'].value
name = ws['B2'].value
description = ws['C2'].value

# открытие документа Word и добавление данных
doc = Document('template.docx')
doc.add_paragraph('Дата: ' + str(date))
doc.add_paragraph('Имя: ' + str(name))
doc.add_paragraph('Описание: ' + str(description))

# сохранение документа
doc.save('output.docx')
