from tkinter import Tk, Label, Text, Button

import pandas as pd
from docx import Document
from docx.shared import Inches

def save_doc():
    # Создание нового документа Word
    doc = Document()

    # Добавление заголовка
    doc.add_heading('Тестовый документ', 0)

    # Загрузка текста из файла Excel
    df = pd.read_excel('test.xlsx')

    # Преобразование данных в список строк
    text_list = df['Текст'].tolist()

    # Добавление текста в документ
    for text in text_list:
        doc.add_paragraph(text)

    # Сохранение документа
    doc.save('test.docx')

# Создание окна Tkinter
root = Tk()

# Добавление элементов управления
label = Label(root, text="Введите текст:")
label.pack()

text_box = Text(root)
text_box.pack()

button = Button(root, text="Сохранить", command=save_doc)
button.pack()

# Запуск главного цикла окна Tkinter
root.mainloop()
